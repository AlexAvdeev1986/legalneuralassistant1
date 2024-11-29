# legal_app/views.py
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_http_methods
from django.core.exceptions import ValidationError
from django.views.decorators.csrf import ensure_csrf_cookie
from django.core.validators import validate_ipv46_address
from django_ratelimit.decorators import ratelimit
from django.shortcuts import get_object_or_404
from .forms import LegalQuestionForm, DocumentGeneratorForm
from .models import LegalQuestion, Document  # Ensure Document is imported
from .utils import LegalAI
import logging
import json
from functools import wraps
from docx import Document as DocxDocument
from typing import Dict, Any

# Initialize LegalAI instance and logger
legal_ai = LegalAI()
logger = logging.getLogger(__name__)

def format_russian_response(data: Dict[str, Any], status: int = 200) -> JsonResponse:
    """Formats a JSON response supporting Russian language."""
    return JsonResponse(data, json_dumps_params={'ensure_ascii': False}, status=status)

@ensure_csrf_cookie
def home(request):
    """Renders the home page."""
    return render(request, 'legal_app/home.html', {
        'title': 'Главная страница',
        'description': 'Юридический помощник с искусственным интеллектом'
    })

def handle_errors(view_func):
    """Decorator for handling common errors in views with Russian messages."""
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        try:
            return view_func(request, *args, **kwargs)
        except ValidationError as e:
            logger.warning(f"Validation error: {str(e)}")
            return format_russian_response({'status': 'error', 'message': 'Ошибка валидации данных.'}, 400)
        except json.JSONDecodeError as e:
            logger.warning(f"JSON decode error: {str(e)}")
            return format_russian_response({'status': 'error', 'message': 'Некорректный формат JSON.'}, 400)
        except Exception as e:
            logger.error(f"Internal error: {str(e)}", exc_info=True)
            return format_russian_response({'status': 'error', 'message': 'Произошла внутренняя ошибка сервера.'}, 500)
    return wrapper

def get_client_ip(request) -> str:
    """Retrieve and validate the client IP address."""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    ip = x_forwarded_for.split(',')[0] if x_forwarded_for else request.META.get('REMOTE_ADDR')
    
    try:
        validate_ipv46_address(ip)
        return ip
    except ValidationError:
        return 'unknown'

@login_required
@require_http_methods(["GET", "POST"])
@ratelimit(key='user', rate='10/m', group='legal_chat')
@handle_errors
def chat(request):
    """Handle legal chat interactions and generate responses."""
    if request.method == 'POST':
        if getattr(request, 'limited', False):
            return format_russian_response({'status': 'error', 'message': 'Превышен лимит запросов. Подождите минуту.'}, 429)
        
        form = LegalQuestionForm(request.POST)
        if form.is_valid():
            question = form.cleaned_data['question']
            user_ip = get_client_ip(request)
            
            response = legal_ai.get_legal_response(question)
            if response.get('status') == 'success':
                legal_question = LegalQuestion.objects.create(
                    user=request.user,
                    question=question,
                    answer=response['answer'],
                    category=response['category'],
                    ip_address=user_ip,
                    status='answered'
                )
                return format_russian_response({
                    'status': 'success',
                    'answer': response['answer'],
                    'category': response['category'],
                    'question_id': legal_question.id
                })
            else:
                return format_russian_response({'status': 'error', 'message': 'Не удалось получить ответ от AI.'}, 500)
        
        return format_russian_response({'status': 'error', 'errors': form.errors}, 400)

    # GET request: display form and question history
    questions = LegalQuestion.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'legal_app/chat.html', {'form': LegalQuestionForm(), 'questions': questions})

@login_required
@require_http_methods(["GET", "POST"])
@handle_errors
def document_generator(request):
    """Generate legal documents dynamically."""
    if request.method == 'POST':
        form = DocumentGeneratorForm(request.POST)
        if form.is_valid():
            try:
                doc_type = form.cleaned_data['document_type']
                title = form.cleaned_data['title']
                context = form.cleaned_data['context']  # Direct plain text context
                
                # Generate document using LegalAI
                generated_content = legal_ai.generate_document(doc_type, context)
                
                # Create Word document
                document = DocxDocument()
                document.add_heading(title, 0)
                for paragraph in generated_content.split("\n"):
                    document.add_paragraph(paragraph)
                
                # Prepare document for download
                response = HttpResponse(
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
                document.save(response)
                
                logger.info(f"Document '{title}' generated successfully.")
                return response
            
            except Exception as e:
                logger.error(f"Error generating document: {str(e)}", exc_info=True)
                return format_russian_response({'status': 'error', 'message': 'Ошибка при генерации документа.'}, 500)
        
        return format_russian_response({'status': 'error', 'errors': form.errors}, 400)

    # GET request: display form and document history
    documents = Document.objects.filter(user=request.user).order_by('-created_at')  # Ensure Document is defined
    return render(request, 'legal_app/document_generator.html', {
        'form': DocumentGeneratorForm(),
        'documents': documents  # Ensure 'documents' context is passed for history display
    })

def download_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    response = HttpResponse(document.file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{document.title}.docx"'
    return response
